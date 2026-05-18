"""从 PDF/DOCX 解析为 LangChain Document 列表（含扫描件 OCR 回退）。"""
from __future__ import annotations

import io
import os
from typing import TYPE_CHECKING

from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document

from app.config.config import KNOWLEDGE_BASE_CONFIG
from app.rag.text_cleaning import clean_text

if TYPE_CHECKING:
    pass


def _ocr_image_bytes(png_bytes: bytes) -> str:
    try:
        import numpy as np
        from PIL import Image
        from rapidocr_onnxruntime import RapidOCR
    except ImportError as e:
        raise RuntimeError(f"OCR 依赖未安装: {e}") from e

    engine = RapidOCR()
    img = Image.open(io.BytesIO(png_bytes))
    arr = np.array(img)
    raw = engine(arr)
    if isinstance(raw, tuple) and len(raw) >= 1:
        ocr_out = raw[0]
    else:
        ocr_out = raw
    if not ocr_out:
        return ""
    lines = []
    for item in ocr_out:
        if len(item) >= 2 and item[1]:
            lines.append(str(item[1]))
    return "\n".join(lines)


def _pdf_with_pymupdf_and_ocr(path: str, doc_id: str) -> list[Document]:
    import fitz

    ocr_enabled = KNOWLEDGE_BASE_CONFIG["ocr_enabled"]
    min_chars = KNOWLEDGE_BASE_CONFIG["ocr_min_chars_per_page"]
    out: list[Document] = []
    doc = fitz.open(path)
    try:
        for page_index in range(len(doc)):
            page = doc[page_index]
            raw = page.get_text() or ""
            text = clean_text(raw.strip())
            use_ocr = ocr_enabled and len(text) < min_chars
            if use_ocr:
                try:
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    png = pix.tobytes("png")
                    ocr_text = _ocr_image_bytes(png)
                    text = clean_text(ocr_text.strip()) if ocr_text else text
                except Exception:
                    if not text:
                        text = ""
            meta = {
                "source": path,
                "doc_id": doc_id,
                "page": page_index,
            }
            out.append(Document(page_content=text or " ", metadata=meta))
    finally:
        doc.close()
    return out


def _pdf_with_pypdf(path: str, doc_id: str) -> list[Document]:
    loader = PyPDFLoader(path)
    docs = loader.load()
    for i, d in enumerate(docs):
        d.metadata["source"] = path
        d.metadata["doc_id"] = doc_id
        if "page" not in d.metadata:
            d.metadata["page"] = d.metadata.get("page", i)
        d.page_content = clean_text(d.page_content or "")
    return docs


def _load_pdf(path: str, doc_id: str) -> list[Document]:
    try:
        docs = _pdf_with_pymupdf_and_ocr(path, doc_id)
        if docs and all(len((d.page_content or "").strip()) < 3 for d in docs):
            return _pdf_with_pypdf(path, doc_id)
        return docs
    except Exception:
        return _pdf_with_pypdf(path, doc_id)


def _load_docx(path: str, doc_id: str) -> list[Document]:
    from docx import Document as DocxDocument

    d = DocxDocument(path)
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    content = clean_text("\n".join(parts))
    return [
        Document(
            page_content=content or " ",
            metadata={"source": path, "doc_id": doc_id, "page": 0},
        )
    ]


def load_documents_from_file(file_path: str, doc_id: str) -> list[Document]:
    """按路径加载单个文件为 Document 列表（未做 chunk 切分）。"""
    if not os.path.isfile(file_path):
        return []
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return _load_pdf(file_path, doc_id)
    if ext == ".docx":
        return _load_docx(file_path, doc_id)
    return []


def joined_preview_text(documents: list[Document], max_chars: int = 500_000) -> str:
    text = "\n\n".join(d.page_content for d in documents if d.page_content)
    if len(text) > max_chars:
        return text[:max_chars] + "\n\n...[truncated]"
    return text
